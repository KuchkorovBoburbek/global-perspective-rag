"""Helpers to export analysis results as Markdown and DOCX for download."""

from __future__ import annotations

import io
from typing import List

from docx import Document as DocxDocument
from docx.shared import Pt
from langchain_core.documents import Document


def build_markdown_report(title: str, body_markdown: str) -> str:
    header = f"# {title}\n\n"
    return header + body_markdown.strip() + "\n"


def markdown_report_to_bytes(title: str, body_markdown: str) -> bytes:
    return build_markdown_report(title, body_markdown).encode("utf-8")


def citation_table_rows(docs: List[Document]) -> List[dict]:
    rows = []
    seen = set()
    for doc in docs:
        meta = doc.metadata
        key = (meta.get("source_label"), meta.get("filename"), meta.get("page"), meta.get("chunk_id"))
        if key in seen:
            continue
        seen.add(key)
        rows.append(
            {
                "source_label": meta.get("source_label", "-"),
                "filename": meta.get("filename", "-"),
                "page": meta.get("page", "-"),
                "chunk_id": meta.get("chunk_id", "-"),
                "language": meta.get("language", "-"),
            }
        )
    return rows


def executive_report_to_docx_bytes(
    title: str,
    body_markdown: str,
    docs: List[Document],
) -> bytes:
    """Render an executive report (markdown text) plus a citation table into
    a simple .docx and return the raw bytes for st.download_button."""
    doc = DocxDocument()

    doc.add_heading(title, level=1)

    for line in body_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    rows = citation_table_rows(docs)
    if rows:
        doc.add_heading("근거 목록 (Evidence List)", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Source"
        hdr[1].text = "Filename"
        hdr[2].text = "Page"
        hdr[3].text = "Chunk ID"
        hdr[4].text = "Language"
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row["source_label"])
            cells[1].text = str(row["filename"])
            cells[2].text = str(row["page"])
            cells[3].text = str(row["chunk_id"])
            cells[4].text = str(row["language"])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
